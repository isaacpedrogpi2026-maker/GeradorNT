import streamlit as st
import pandas as pd
import numpy as np
import re
import requests
from io import BytesIO
import openpyxl
from datetime import datetime
import unicodedata
from pathlib import Path
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse, urljoin

st.set_page_config(page_title="Gerador de Nota Técnica", layout="wide")

try:
    import db_supabase
except ImportError:
    db_supabase = None


# Define status order mapping
STATUS_ORDER = ['CONCLUÍDO', 'EM ANDAMENTO', 'A INICIAR', 'EM LICITAÇÃO', 'A LICITAR', 'A FAZER']

# Status mapping for normalization
STATUS_MAPPING = {
    'Concluído': 'CONCLUÍDO',
    'Concluida': 'CONCLUÍDO',
    'Em andamento': 'EM ANDAMENTO',
    'Em Andamento': 'EM ANDAMENTO',
    'A iniciar': 'A INICIAR',
    'A Iniciar': 'A INICIAR',
    'Em licitação': 'EM LICITAÇÃO',
    'Em Licitação': 'EM LICITAÇÃO',
    'A licitar': 'A LICITAR',
    'A Licitar': 'A LICITAR',
    'Projeto em Elaboração': 'A LICITAR',
    'A elaborar projeto': 'A LICITAR',
    'Vamos fazer': 'A FAZER',
    'A fazer': 'A FAZER'
}

def load_excel_file(uploaded_file):
    """Load all sheets from an uploaded Excel file."""
    if uploaded_file is None:
        return {}

    if hasattr(uploaded_file, 'read'):
        raw_data = uploaded_file.read()
    else:
        raw_data = uploaded_file

    if isinstance(raw_data, memoryview):
        raw_data = raw_data.tobytes()

    if not isinstance(raw_data, (bytes, bytearray)):
        raise ValueError("O arquivo enviado não foi lido como bytes.")

    raw_bytes = bytes(raw_data)
    if not raw_bytes:
        raise ValueError("O arquivo enviado está vazio.")

    filename = getattr(uploaded_file, 'name', '') or ''
    file_ext = Path(filename).suffix.lower()

    def detect_engine(raw_bytes, extension):
        if extension in {'.xlsx', '.xlsm', '.xltx', '.xltm'}:
            return 'openpyxl'
        if extension in {'.xls', '.xlt'}:
            return 'xlrd'

        if raw_bytes.startswith(b'PK\x03\x04'):
            return 'openpyxl'
        if raw_bytes.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
            return 'xlrd'

        return None

    def looks_like_zip(raw_bytes):
        return raw_bytes.startswith(b'PK\x03\x04')

    def looks_like_ole(raw_bytes):
        return raw_bytes.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1')

    engine = detect_engine(raw_bytes, file_ext)
    fallback_engines = []
    if engine:
        fallback_engines.append(engine)
    fallback_engines.extend(['openpyxl', 'xlrd'])

    if not looks_like_zip(raw_bytes) and not looks_like_ole(raw_bytes) and file_ext not in {'.xlsx', '.xlsm', '.xltx', '.xltm', '.xls', '.xlt'}:
        raise ValueError(
            "O arquivo não parece ser uma planilha Excel válida. Envie um arquivo .xlsx ou .xls, não um PDF, CSV ou outro formato."
        )

    excel_data = {}
    last_error = None
    for candidate_engine in dict.fromkeys(fallback_engines):
        try:
            buffer = BytesIO(raw_bytes)
            xls = pd.ExcelFile(buffer, engine=candidate_engine)
            for sheet_name in xls.sheet_names:
                sheet_buffer = BytesIO(raw_bytes)
                excel_data[sheet_name] = pd.read_excel(sheet_buffer, sheet_name=sheet_name, engine=candidate_engine)
            return excel_data
        except Exception as exc:
            last_error = exc
            excel_data = {}

    try:
        buffer = BytesIO(raw_bytes)
        xls = pd.ExcelFile(buffer)
        for sheet_name in xls.sheet_names:
            sheet_buffer = BytesIO(raw_bytes)
            excel_data[sheet_name] = pd.read_excel(sheet_buffer, sheet_name=sheet_name)
        return excel_data
    except Exception as exc:
        last_error = exc

    if last_error is not None:
        raise ValueError(
            "Não foi possível ler o arquivo Excel. Envie um arquivo .xlsx ou .xls válido. "
            f"Detalhes: {last_error}"
        ) from last_error
    raise ValueError("Não foi possível ler o arquivo Excel.")


def normalize_status(status):
    """Normalize status values according to mapping"""
    if pd.isna(status):
        return 'A FAZER'
    status_str = str(status).strip()
    return STATUS_MAPPING.get(status_str, status_str.upper())


def get_watermark_image_path():
    parent = Path(__file__).resolve().parent
    # Prefer specific well-known filenames, then fallback to heuristics
    try:
        import unicodedata as _ud
    except Exception:
        _ud = None

    def _norm(name):
        n = name.lower()
        if _ud:
            n = _ud.normalize('NFKD', n)
            n = ''.join(c for c in n if not (_ud.category(c).startswith('M')))
        return n

    # exact known names
    candidates_to_check = [
        'Base NT nítido compesa.png',
        'Base NT nitido compesa.png',
        'Base NT nítido.png',
        'base nt compesa.png',
        'Base NT 30062026.png'
    ]
    for name in candidates_to_check:
        p = parent / name
        if p.exists():
            return p

    # heuristic: look for png files containing base + nt + compesa
    pngs = list(parent.glob('*.png'))
    for p in pngs:
        n = _norm(p.name)
        if 'base' in n and 'nt' in n and 'compesa' in n:
            return p

    # fallback: base + nt
    for p in pngs:
        n = _norm(p.name)
        if 'base' in n and 'nt' in n:
            return p

    # fallback: any png containing 'marca' and 'agua'
    for p in pngs:
        n = _norm(p.name)
        if 'marca' in n and 'agua' in n:
            return p

    # last resort: any png in directory
    return pngs[0] if pngs else None


def create_watermark_image_bytes(opacity=0.15):
    watermark_path = get_watermark_image_path()
    if watermark_path is None:
        return None
    try:
        from PIL import Image
    except ImportError:
        return None

    with Image.open(watermark_path) as img:
        if img.mode != 'RGBA':
            img = img.convert('RGBA')
        # apply requested opacity to alpha channel
        if opacity is not None and 0 <= opacity < 1.0:
            alpha = img.split()[3].point(lambda p: int(p * opacity))
            img.putalpha(alpha)
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        return buffer


def download_excel_file(url):
    """Download an Excel file from a cloud link."""
    if not url or not isinstance(url, str):
        raise ValueError("URL inválida para download do Excel.")

    url = url.strip()
    parsed = urlparse(url)
    netloc = parsed.netloc.lower()
    query = parse_qs(parsed.query, keep_blank_values=True)

    if 'drive.google.com' in url and '/file/d/' in url:
        file_id = url.split('/file/d/')[1].split('/')[0]
        url = f'https://drive.google.com/uc?export=download&id={file_id}'
    elif 'docs.google.com/spreadsheets/d/' in url:
        file_id = url.split('/spreadsheets/d/')[1].split('/')[0]
        url = f'https://docs.google.com/spreadsheets/d/{file_id}/export?format=xlsx'
    elif 'onedrive.live.com' in netloc or 'sharepoint.com' in netloc or '1drv.ms' in netloc:
        # For Microsoft cloud sharing links, force the file to download.
        if 'download' not in query:
            query['download'] = ['1']
        if query.get('action') == ['default']:
            query.pop('action', None)
        new_query = urlencode(query, doseq=True)
        url = urlunparse((parsed.scheme, parsed.netloc, parsed.path, parsed.params, new_query, parsed.fragment))

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:118.0) Gecko/20100101 Firefox/118.0'
    }

    def fetch_content(request_url, timeout, stream=False):
        response = requests.get(
            request_url,
            headers=headers,
            allow_redirects=True,
            timeout=timeout,
            stream=stream,
        )
        response.raise_for_status()

        if stream:
            buffer = BytesIO()
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    buffer.write(chunk)
            buffer.seek(0)
            return buffer, response.headers

        return BytesIO(response.content), response.headers

    try:
        raw_buffer, headers_response = fetch_content(url, timeout=(10, 120), stream=False)
    except requests.exceptions.ReadTimeout:
        try:
            raw_buffer, headers_response = fetch_content(url, timeout=(10, 300), stream=True)
        except requests.exceptions.ReadTimeout as exc:
            raise ValueError(
                'O download demorou mais do que o esperado. Tente novamente mais tarde ou baixe o arquivo localmente e envie pelo carregamento direto.'
            ) from exc
    except requests.exceptions.HTTPError as exc:
        if exc.response is not None and exc.response.status_code == 403:
            raise ValueError(
                'O link exige autenticação ou permissões. Baixe o arquivo manualmente e envie pelo upload local, ou use um link público direto de download.'
            ) from exc
        raise

    content_type = headers_response.get('Content-Type', '').lower()
    if 'html' in content_type and b'<html' in raw_buffer.getvalue()[:400].lower():
        # The URL returned an HTML page (common for SharePoint/OneDrive redirect pages).
        # Try to find a direct download link inside the HTML and retry.
        try:
            html = raw_buffer.getvalue().decode('utf-8', errors='ignore')
            # look for href that likely points to a download/export or a workbook
            import re
            m = re.search(r'href=["\']([^"\']+(?:download|export|GetFile|_layouts/15/download.aspx|\.xlsx|\.xls)[^"\']*)["\']', html, re.IGNORECASE)
            if m:
                candidate = m.group(1)
                candidate_url = urljoin(url, candidate)
                try:
                    raw_buffer2, headers_response2 = fetch_content(candidate_url, timeout=(10, 120), stream=False)
                    # if this looks like a file, return it
                    ctype2 = headers_response2.get('Content-Type', '').lower()
                    if 'html' not in ctype2:
                        return raw_buffer2
                except Exception:
                    pass
        except Exception:
            pass

        raise ValueError('O link retornou HTML em vez de um arquivo Excel. Verifique se o link é direto para download ou forneça um link de compartilhamento direto com download.')

    return raw_buffer


def normalize_text(value):
    if pd.isna(value):
        return ''
    text = str(value).strip()
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('ASCII')
    return text.upper()


def normalize_column_name(col):
    normalized = normalize_text(col)
    normalized = re.sub(r'[^A-Z0-9 ]+', ' ', normalized)
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    return normalized


def find_column(df, candidates):
    if df is None or df.empty:
        return None

    candidate_norms = [normalize_column_name(c) for c in candidates]
    columns_norm = {col: normalize_column_name(col) for col in df.columns}

    for candidate_norm in candidate_norms:
        for col, col_norm in columns_norm.items():
            if col_norm == candidate_norm:
                return col

    for candidate_norm in candidate_norms:
        for col, col_norm in columns_norm.items():
            if candidate_norm and candidate_norm in col_norm:
                return col
            if col_norm and col_norm in candidate_norm:
                return col

    return None


def find_best_investment_column(df):
    if df is None or df.empty:
        return None

    # Prefer the per-municipio value when available, then the total field.
    preferred = find_column(df, [
        'Total do Investimento R$ por município',
        'Total do Investimento por município',
        'Total do Investimento por municipio',
        'Total do Investimento no município',
        'Total do Investimento no município (R$)',
        'Investimento no município',
        'Investimento no municipio',
        'Investimento no município (R$)',
        'Investimento no municipio (R$)'
    ])
    if preferred:
        return preferred

    fallback = find_column(df, [
        'Total do Investimento R$',
        'Total do Investimento',
        'Total Investimento',
        'Investimento Total'
    ])
    if fallback:
        return fallback

    # If there is any field with investment in the normalized name, choose the first one.
    for col in df.columns:
        normalized = normalize_column_name(col)
        if 'TOTALDOINVESTIMENTO' in normalized.replace(' ', '') or 'INVESTIMENTO' in normalized:
            return col

    return None


def normalize_dict_key(key):
    return normalize_column_name(key) if key is not None else None


def get_first_key(data, keys, default=None):
    if data is None or not isinstance(data, dict):
        return default

    normalized_data = {normalize_dict_key(k): v for k, v in data.items()}
    for key in keys:
        normalized_key = normalize_dict_key(key)
        if normalized_key in normalized_data:
            return normalized_data[normalized_key]
    return default


def get_numeric_value(data, keys, default=0):
    value = get_first_key(data, keys, default)
    numeric = to_numeric_series(pd.Series([value])).iloc[0]
    return default if pd.isna(numeric) else numeric


def get_investimento_municipio(data, default=0):
    return get_numeric_value(
        data,
        [
            'Total do Investimento R$ por município',
            'Total do Investimento por município',
            'Total do Investimento por municipio',
            'Total do Investimento no município',
            'Total do Investimento no município (R$)',
            'Investimento no município',
            'Investimento no municipio',
            'Investimento no município (R$)',
            'Investimento no municipio (R$)'
        ],
        default,
    )


def is_municipio_column(col_name):
    normalized = normalize_column_name(col_name)
    candidates = ['MUNICIPIO', 'MUNICIPIO PRINCIPAL', 'CIDADE', 'LOCALIDADE', 'MUNICIPALITY']
    return any(candidate in normalized for candidate in candidates)


def find_municipio_column(df, sheet_name=None):
    """Find a municipality-like column name in the dataframe."""
    if df is None or df.empty:
        return None

    cols = find_municipio_columns(df, sheet_name)
    return cols[0] if cols else None


def find_municipio_columns(df, sheet_name=None):
    """Find all municipality-like columns in the dataframe."""
    if df is None or df.empty:
        return []

    return [col for col in df.columns if is_municipio_column(col)]


def find_sheet_name(dataframes, candidates, required_columns=None):
    """Find a worksheet by name or by the presence of expected columns."""
    if not dataframes:
        return None

    required_normalized = {normalize_column_name(col) for col in (required_columns or [])}

    for sheet_name in dataframes.keys():
        df = dataframes.get(sheet_name)
        if df is None or df.empty:
            continue

        normalized_sheet = normalize_column_name(sheet_name)
        if any(normalized_sheet == normalize_column_name(candidate) for candidate in candidates):
            if required_columns is None or any(normalize_column_name(col) in required_normalized for col in df.columns):
                return sheet_name

    for sheet_name in dataframes.keys():
        df = dataframes.get(sheet_name)
        if df is None or df.empty:
            continue

        if required_columns is None:
            return sheet_name

        if any(normalize_column_name(col) in required_normalized for col in df.columns):
            return sheet_name

    return None


def collect_municipios(dataframes):
    """Collect unique municipalities from the segregation sheet using the Municipio Principal column."""
    municipio_values = set()

    segreg_df = None
    for sheet_name in ['Segreg por município', 'Segreg por municipio', 'Investimentos', 'Investimento']:
        candidate = dataframes.get(sheet_name)
        if candidate is not None and not getattr(candidate, 'empty', True):
            segreg_df = candidate
            break

    if segreg_df is not None and not getattr(segreg_df, 'empty', True):
        municipio_col = find_column(segreg_df, ['Município Principal', 'Municipio Principal', 'Município', 'Municipio'])
        if municipio_col:
            values = segreg_df[municipio_col].dropna().astype(str).str.strip()
            municipio_values.update([v for v in values if v])

    if municipio_values:
        return sorted(municipio_values, key=lambda x: str(x))

    for sheet in ['Investimentos', 'Investimento', 'Volumes por município', 'Volumes por municipio', 'Abastecimento', 'População por município', 'Populacao por município']:
        df = dataframes.get(sheet)
        if df is None or df.empty:
            continue

        for col in find_municipio_columns(df, sheet):
            values = df[col].dropna().astype(str).str.strip()
            municipio_values.update([v for v in values if v])

    return sorted(municipio_values, key=lambda x: str(x))


def filter_by_municipio(df, municipio, sheet_name=None):
    """Filter dataframe by selected municipality using the municipality column from the segregation sheet."""
    if df is None or df.empty or municipio is None:
        return df

    municipio_col = None
    if sheet_name in ['Segreg por município', 'Segreg por municipio', 'Investimentos', 'Investimento']:
        municipio_col = find_column(df, ['Município Principal', 'Municipio Principal', 'Município', 'Municipio'])

    if municipio_col is None:
        municipio_cols = find_municipio_columns(df, sheet_name)
        if not municipio_cols:
            return df
        municipio_col = municipio_cols[0]

    target = normalize_text(municipio)
    values = df[municipio_col].apply(normalize_text)
    return df[values == target]


def to_numeric_series(series):
    """Convert a series to numeric values, coercing invalid entries to NaN."""
    if series is None:
        return pd.Series(dtype='float64')

    def parse_value(value):
        if pd.isna(value):
            return np.nan

        if isinstance(value, (int, float)):
            return float(value)

        text = str(value).strip()
        if text == '' or text.lower() in ['nan', 'none', 'na']:
            return np.nan

        text = text.replace('R$', '').replace('\xa0', ' ')
        raw_text = text.lower().strip()

        has_milhoes = bool(re.search(r'milh(ões|oes|ao)s?$', raw_text))
        has_mil = bool(re.search(r'\bmil\b', raw_text))
        has_k = raw_text.endswith('k')

        if has_milhoes:
            raw_text = re.sub(r'\s*milh(ões|oes|ao)s?\s*$', '', raw_text)
        elif has_mil:
            raw_text = re.sub(r'\s*\bmil\b\s*$', '', raw_text)
        elif has_k:
            raw_text = raw_text[:-1]

        raw_text = raw_text.strip()
        raw_text = re.sub(r'[^0-9,.-]', '', raw_text)

        if raw_text.count('.') > 1 and raw_text.count(',') == 0:
            raw_text = raw_text.replace('.', '')
        elif raw_text.count('.') > 0 and raw_text.count(',') > 0:
            if raw_text.rfind(',') > raw_text.rfind('.'):
                raw_text = raw_text.replace('.', '')
                raw_text = raw_text.replace(',', '.')
            else:
                raw_text = raw_text.replace(',', '')
        elif raw_text.count(',') > 0 and raw_text.count('.') == 0:
            raw_text = raw_text.replace(',', '.')

        if raw_text == '' or raw_text in ['.', ',', '-']:
            return np.nan

        try:
            base_value = float(raw_text)
        except ValueError:
            return np.nan

        if has_milhoes and base_value <= 9999:
            return base_value * 1_000_000
        if (has_mil or has_k) and base_value <= 9999:
            return base_value * 1_000
        return base_value

    return pd.Series(series.apply(parse_value), index=series.index, dtype='float64')

def get_first_non_empty_value(df, column_candidates, default='[Informação não disponível]'):
    if df is None or df.empty:
        return default

    col = find_column(df, column_candidates)
    if col is None or col not in df.columns:
        return default

    for value in df[col].dropna():
        text = str(value).strip()
        if text not in ['', 'nan', 'None']:
            return value

    return default


def normalize_missing_value(value, default='[Informação não disponível]'):
    if pd.isna(value):
        return default
    if isinstance(value, str):
        text = value.strip()
        if text == '' or text.lower() in ['n/a', 'na', 'nan', 'none']:
            return default
        return value
    return value


def has_effective_value(value):
    normalized = normalize_missing_value(value, default='').strip()
    return normalized != ''


def format_population_value(value):
    if pd.isna(value):
        return '[Informação não disponível]'

    if isinstance(value, str):
        text = value.strip()
        if text == '' or text.lower() in ['n/a', 'na', 'nan', 'none']:
            return '[Informação não disponível]'

    numeric = to_numeric_series(pd.Series([value])).iloc[0]
    if pd.isna(numeric):
        return str(value)

    if abs(float(numeric)) < 1000:
        if float(numeric).is_integer():
            return format_integer_pt_br(int(numeric))
        return format_plain_number_pt_br(numeric)

    return format_plain_number_pt_br(numeric)


def normalize_report_flag(value):
    if pd.isna(value):
        return None

    text = normalize_text(value)
    if text in {'SIM', 'S', 'YES', 'Y', 'TRUE', 'T', '1'}:
        return True
    if text in {'NAO', 'NAO', 'NÃO', 'NO', 'N', 'FALSE', 'F', '0'}:
        return False
    return None


def filter_by_informar_no_relatorio(df, option='Todos'):
    if df is None or df.empty or option in [None, 'Todos']:
        return df

    col = find_column(df, ['Informar no Relatório', 'Informar no Relatorio', 'Informar no relatorio'])
    if col is None or col not in df.columns:
        return df

    if option == 'Sim':
        values = df[col].apply(normalize_report_flag)
        return df[values == True]

    if option == 'Não':
        values = df[col].apply(normalize_report_flag)
        return df[values == False]

    return df


def prepare_obras_for_display(df_segreg):
    if df_segreg is None:
        return pd.DataFrame()

    df = df_segreg.copy()
    if df.empty:
        return df

    status_col = find_column(df, ['Status', 'STATUS'])
    if status_col and status_col in df.columns:
        df['Status_Normalizado'] = df[status_col].apply(normalize_status)
    else:
        df['Status_Normalizado'] = pd.Series(['A FAZER'] * len(df), index=df.index)

    valor_col = find_best_investment_column(df)
    if valor_col and valor_col in df.columns:
        df['__valor_investimento__'] = to_numeric_series(df[valor_col])
    else:
        df['__valor_investimento__'] = pd.Series([0.0] * len(df), index=df.index)

    fonte_col = find_column(df, ['Fonte de Recurso', 'Fonte de recurso'])
    if fonte_col and fonte_col in df.columns:
        df['__exibir__'] = df[fonte_col].apply(lambda value: normalize_text(value) != 'A CAPTAR')
    else:
        df['__exibir__'] = pd.Series([True] * len(df), index=df.index)

    return df


def extract_first_numeric_from_column(df, column_candidates, default='[Informação não disponível]'):
    if df is None or df.empty:
        return default

    col = find_column(df, column_candidates)
    if col is None or col not in df.columns:
        return default

    values = to_numeric_series(df[col]).dropna()
    if values.empty:
        return default

    return format_plain_number_pt_br(values.iloc[0])


def extract_dados_gerais(df_segreg, df_volumes, df_abastecimento, df_populacao, df_tv_tspe=None):
    """Extract general data from the dataframes"""
    dados_gerais = {}
    
    # 1 - Sistema de abastecimento (from Abastecimento)
    dados_gerais['sistema_abastecimento'] = get_first_non_empty_value(df_abastecimento, ['Sistema de Abastecimento', 'Sistema de abastecimento', 'Sistema'])
    
    # 2 - Produção total de água (from Abastecimento)
    vazao_col = find_column(df_abastecimento, ['Vazão (L/s)', 'Vazao (L/s)'])
    if vazao_col and vazao_col in df_abastecimento.columns:
        vazao = to_numeric_series(df_abastecimento[vazao_col])
        dados_gerais['producao_total_agua'] = f"{vazao.sum():.0f} L/s"
    else:
        dados_gerais['producao_total_agua'] = '[Informação não disponível]'
    
    # 3 - Calendário médio de abastecimento (from Abastecimento)
    dados_gerais['calendario_medio'] = get_first_non_empty_value(df_abastecimento, ['Calendário Médio Atual', 'Calendario Medio Atual', 'Calendario Medio', 'Calendário Atual'])
    
    # 4 - Calendário Médio Pós Obras (from Abastecimento)
    dados_gerais['calendario_pos_obras'] = get_first_non_empty_value(df_abastecimento, ['Calendário Médio Pós Obras', 'Calendario Medio Pos Obras', 'Calendario Medio Pos', 'Calendário Pós Obras'])
    
    # 5 - Famílias beneficiadas (from TV+TSPE sheet, column TV + TSPE (Econ))
    dados_gerais['familias_beneficiadas'] = extract_first_numeric_from_column(
        df_tv_tspe,
        ['TV + TSPE (Econ)', 'TV + TSPE', 'TV TSPE']
    )
    
    # 6 - População (from Abastecimento sheet, column População Total Residente SINISA 2024)
    dados_gerais['populacao'] = extract_first_numeric_from_column(
        df_abastecimento,
        ['População Total Residente SINISA 2024', 'Populacao Total Residente SINISA 2024']
    )
    
    return dados_gerais

def extract_dados_obras(df_segreg):
    """Extract obra data with status grouping"""
    df = prepare_obras_for_display(df_segreg)

    status_summary = {}
    for status in STATUS_ORDER:
        filtered = df[(df['Status_Normalizado'] == status) & df['__exibir__']].copy()
        if not filtered.empty:
            filtered = filtered.sort_values('__valor_investimento__', ascending=False)
            status_summary[status] = {
                'quantidade': int(len(filtered)),
                'valor': float(filtered['__valor_investimento__'].sum())
            }
        else:
            status_summary[status] = {
                'quantidade': 0,
                'valor': 0
            }

    return status_summary


def extract_descricao_obras(df_segreg):
    """Extract detailed description of obras"""
    df = prepare_obras_for_display(df_segreg)

    obras_por_status = {}
    for status in STATUS_ORDER:
        obras = df[(df['Status_Normalizado'] == status) & df['__exibir__']].copy()
        if not obras.empty:
            obras = obras.sort_values('__valor_investimento__', ascending=False)
            obras_por_status[status] = obras.drop(columns=['Status_Normalizado', '__valor_investimento__', '__exibir__']).to_dict('records')

    return obras_por_status

def format_currency(value):
    """Format value as Brazilian currency"""
    if pd.isna(value) or value == 0:
        return "R$ 0,00"
    return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def format_currency_rounded(value):
    """Format a currency value in pt_BR with thousands/millions/billions rounding."""
    if pd.isna(value) or value == 0:
        return "R$ 0"

    n = float(value)
    abs_n = abs(n)

    def format_with_unit(base_value, unit_single, unit_plural):
        if base_value.is_integer():
            formatted = f"{int(base_value):,}"
        else:
            formatted = f"{base_value:,.1f}"
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
        if 1 <= abs(base_value) < 2:
            return f"R$ {formatted} {unit_single}"
        return f"R$ {formatted} {unit_plural}"

    if abs_n >= 1_000_000_000:
        billions = n / 1_000_000_000
        return format_with_unit(billions, 'bilhão', 'bilhões')
    if abs_n >= 1_000_000:
        millions = n / 1_000_000
        return format_with_unit(millions, 'milhão', 'milhões')
    if abs_n >= 1_000:
        thousands = n / 1_000
        return format_with_unit(thousands, 'mil', 'mil')

    return format_currency(n)


def format_plain_number_pt_br(value):
    """Format summary values in pt_BR without the currency symbol."""
    if pd.isna(value) or value == 0:
        return "0"

    n = float(value)
    abs_n = abs(n)

    if abs_n >= 1_000_000:
        millions = n / 1_000_000
        if millions.is_integer():
            formatted = f"{int(millions):,}"
        else:
            formatted = f"{millions:,.1f}"
        unit = 'milhão' if 1 <= abs(millions) < 2 else 'milhões'
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
        return f"{formatted} {unit}"
    if abs_n >= 1_000:
        thousands = n / 1_000
        if thousands.is_integer():
            formatted = f"{int(thousands):,}"
        else:
            formatted = f"{thousands:,.1f}"
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
        return f"{formatted} mil"

    formatted = f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return formatted


def format_integer_pt_br(value):
    """Format integer values in pt_BR without decimals."""
    if pd.isna(value) or value in ['', 'nan', 'None', None]:
        return ""

    try:
        n = float(value)
    except Exception:
        return str(value).strip()

    if np.isnan(n):
        return ""

    n = int(n)
    formatted = f"{n:,}".replace(",", ".")
    return formatted


def format_date_pt_br(value):
    """Format date-like values to dd/mm/aaaa."""
    if pd.isna(value) or value in ['', 'nan', 'None', None]:
        return ""

    try:
        parsed = pd.to_datetime(value, dayfirst=True, errors='coerce')
    except Exception:
        parsed = pd.NaT

    if isinstance(parsed, pd.Timestamp) and pd.notna(parsed):
        return parsed.strftime('%d/%m/%Y')

    if isinstance(value, str):
        return value.strip()

    return str(value)


def format_month_year_pt_br(value):
    """Format date-like values to Mês/Ano in Portuguese."""
    if pd.isna(value) or value in ['', 'nan', 'None', None]:
        return ""

    try:
        parsed = pd.to_datetime(value, dayfirst=True, errors='coerce')
    except Exception:
        parsed = pd.NaT

    if isinstance(parsed, pd.Timestamp) and pd.notna(parsed):
        month_names = [
            '', 'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
            'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
        ]
        return f"{month_names[parsed.month].capitalize()}/{parsed.year}"

    if isinstance(value, str):
        normalized = value.strip()
        parts = re.split(r'[\/\-\.\s]+', normalized)
        if len(parts) >= 2:
            try:
                year = int(parts[-1])
                month = int(parts[-2])
                if 1 <= month <= 12:
                    month_names = [
                        '', 'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
                        'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
                    ]
                    return f"{month_names[month].capitalize()}/{year}"
            except Exception:
                pass

    return str(value).strip()


def format_month_year_numeric_pt_br(value):
    """Format date-like values as mm/YYYY."""
    if pd.isna(value) or value in ['', 'nan', 'None', None]:
        return ""

    try:
        parsed = pd.to_datetime(value, dayfirst=True, errors='coerce')
    except Exception:
        parsed = pd.NaT

    if isinstance(parsed, pd.Timestamp) and pd.notna(parsed):
        return parsed.strftime('%m/%Y')

    if isinstance(value, str):
        normalized = value.strip()
        parts = re.split(r'[\/\-\.\s]+', normalized)
        if len(parts) >= 2:
            try:
                year = int(parts[-1])
                month = int(parts[-2])
                if 1 <= month <= 12:
                    return f"{month:02d}/{year}"
            except Exception:
                pass
        return normalized

    return str(value).strip()


def build_status_highlight_color(status):
    return {
        'CONCLUÍDO': 'D9EAD3',
        'CONCLUÍDA': 'D9EAD3',
        'EM ANDAMENTO': 'FFF2CC',
        'A INICIAR': 'FCE4D6',
        'A LICITAR': 'F4CCCC',
        'EM LICITAÇÃO': 'F2F2F2',
        'A FAZER': 'F2F2F2',
        'TOTAL': 'F2F2F2'
    }.get(status, 'F2F2F2')


def build_obra_details(status, obra):
    status_display = {
        'CONCLUÍDO': 'CONCLUÍDA',
        'EM ANDAMENTO': 'EM ANDAMENTO',
        'A INICIAR': 'A INICIAR',
        'EM LICITAÇÃO': 'EM LICITAÇÃO',
        'A LICITAR': 'A LICITAR',
        'A FAZER': 'A FAZER'
    }

    nome_iniciativa = normalize_missing_value(get_first_key(obra, ['Nome da Iniciativa', 'Nome da iniciativa'], '[Informação não disponível]'))
    pop_beneficiada = format_population_value(get_first_key(obra, ['População Beneficiada', 'Populacao Beneficiada'], '[Informação não disponível]'))
    municipios = normalize_missing_value(get_first_key(obra, ['Todos os municípios beneficiados', 'Todos os municipios beneficiados'], '[Informação não disponível]'))
    investimento = get_investimento_municipio(obra, 0)
    investimento_total = get_numeric_value(
        obra,
        ['Total do Investimento da iniciativa R$', 'Total do Investimento da iniciativa', 'Total do Investimento R$', 'Total do Investimento R$ por município'],
        investimento,
    )
    fonte = normalize_missing_value(get_first_key(obra, ['Fonte de Recurso', 'Fonte de recurso'], 'Governo de Pernambuco e Compesa'))
    prazo = get_first_key(obra, ['Prazo de Conclusão', 'Prazo de Conclusao'], '')
    data_inicio = get_first_key(obra, ['Data Início', 'Data Inicio'], '')
    proxima_etapa = get_first_key(obra, ['Próxima Etapa', 'Proxima Etapa'], '')
    data_proxima_etapa = get_first_key(obra, ['Data Próxima Etapa', 'Data Proxima Etapa'], '')

    if status == 'CONCLUÍDO' and prazo:
        status_row = ('Conclusão', format_month_year_pt_br(prazo))
    elif status == 'EM ANDAMENTO' and prazo:
        status_row = ('Previsão de término', format_month_year_pt_br(prazo))
    elif status == 'EM LICITAÇÃO':
        status_row = None
        if data_inicio:
            status_row = ('Previsão de início', format_month_year_pt_br(data_inicio))
        elif prazo:
            status_row = ('Previsão de término', format_month_year_pt_br(prazo))
    elif status in ('A INICIAR', 'A LICITAR', 'A FAZER'):
        formatted_proxima = format_date_pt_br(proxima_etapa)
        if status == 'A LICITAR' or status == 'A INICIAR':
            formatted_data = format_month_year_pt_br(data_proxima_etapa)
        else:
            formatted_data = format_date_pt_br(data_proxima_etapa)
        data_label = 'Previsão de emissão de ordem de serviço:' if status in ('A INICIAR', 'A FAZER') else 'Previsão de publicação do edital de licitação:'

        if formatted_proxima and formatted_data:
            status_row = ('Próxima etapa', f"{formatted_proxima}. {data_label} {formatted_data}")
        elif formatted_proxima:
            status_row = ('Próxima etapa', f"{formatted_proxima}.")
        elif formatted_data:
            status_row = ('Próxima etapa', f"{data_label} {formatted_data}")
        else:
            status_row = None
    else:
        status_row = None

    detail_rows = [
        ('População beneficiada', pop_beneficiada),
        ('Municípios beneficiados', municipios),
        ('Investimento no município', f"{format_currency_rounded(investimento)} ({fonte})"),
        ('Total do Investimento', format_currency_rounded(investimento_total))
    ]
    if status_row:
        detail_rows.append(status_row)

    return {
        'status': status_display.get(status, status),
        'status_color': build_status_highlight_color(status),
        'initiative': nome_iniciativa,
        'rows': detail_rows
    }


def build_document_model(selected_obras_by_status, dados_gerais, dados_obras, selected_municipio):
    status_display = {
        'CONCLUÍDO': 'CONCLUÍDA',
        'EM ANDAMENTO': 'EM ANDAMENTO',
        'A INICIAR': 'A INICIAR',
        'EM LICITAÇÃO': 'EM LICITAÇÃO',
        'A LICITAR': 'A LICITAR',
        'A FAZER': 'A FAZER'
    }

    calendario_medio = normalize_missing_value(dados_gerais.get('calendario_medio', '[Informação não disponível]'))
    calendario_pos_obras = normalize_missing_value(dados_gerais.get('calendario_pos_obras', '[Informação não disponível]'))
    if calendario_pos_obras == '[Informação não disponível]':
        calendario_pos_obras = ''
    if calendario_medio and calendario_pos_obras and calendario_medio == calendario_pos_obras:
        calendario_pos_obras = ''

    general = [
        ('Sistemas de abastecimento', normalize_missing_value(dados_gerais.get('sistema_abastecimento', '[Informação não disponível]'))),
        ('Produção total de água', normalize_missing_value(dados_gerais.get('producao_total_agua', '[Informação não disponível]'))),
        ('Calendário médio de abastecimento', calendario_medio),
        ('Calendário médio pós obras', calendario_pos_obras),
        ('Famílias beneficiadas com a Tarifa Social Pernambucana e a Tarifa de Vulneráveis', format_integer_pt_br(dados_gerais.get('familias_beneficiadas', '[Informação não disponível]'))),
        ('População', normalize_missing_value(dados_gerais.get('populacao', '[Informação não disponível]')))
    ]
    general = [row for row in general if has_effective_value(row[1])]

    summary = []
    for row in summarize_obras_for_report(selected_obras_by_status=selected_obras_by_status, dados_obras=dados_obras):
        summary.append({
            'status': status_display.get(row['Status'], row['Status']),
            'quantidade': int(row['Quantidade']),
            'valor': format_currency_rounded(row['Valor (R$)'])
        })

    descricao = []
    for status in STATUS_ORDER:
        obras = selected_obras_by_status.get(status, [])
        if obras:
            descricao.append({
                'status': status_display.get(status, status),
                'details': [build_obra_details(status, obra) for obra in obras]
            })

    generated_at = datetime.now().strftime('%d/%m/%Y')
    return {
        'title': f'Município de {selected_municipio}',
        'general': general,
        'summary': summary,
        'descricao': descricao,
        'generated_at': generated_at,
        'updated_text': f'Atualizada em: {generated_at}'
    }


def create_document_text(selected_obras_by_status, dados_gerais, dados_obras, selected_municipio):
    model = build_document_model(selected_obras_by_status, dados_gerais, dados_obras, selected_municipio)
    lines = [model['title'], '', 'DADOS GERAIS']
    for label, value in model['general']:
        lines.append(f'  • {label}: {value}')
    lines.append('')
    lines.append('DADOS DAS OBRAS')
    for row in model['summary']:
        lines.append(f"  • {row['status']}: Quantidade {row['quantidade']} — Valor (R$): {row['valor']}")
    lines.append('')
    lines.append('DETALHAMENTO DAS OBRAS')
    for group in model['descricao']:
        for details in group['details']:
            lines.append(f"{details['status']} - {details['initiative']}")
            for label, value in details['rows']:
                lines.append(f"  • {label}: {value}")
            lines.append('')
    return '\n'.join(lines).strip()


def create_doc_bytes(document_text):
    return document_text.encode('utf-8-sig')


def create_docx_bytes(model):
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    def clean_value(value):
        return value.replace('R$', '').strip()

    def wrap_text_for_cell(text, max_chars=45):
        words = str(text).split()
        if not words:
            return ['']
        lines = []
        current = words[0]
        for word in words[1:]:
            if len(current) + 1 + len(word) <= max_chars:
                current += ' ' + word
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    def set_cell_text(cell, text, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, bold=False):
        # Put wrapped lines into separate paragraphs inside the cell
        lines = wrap_text_for_cell(text, max_chars=45)
        # assign to first paragraph
        if cell.paragraphs:
            paragraph = cell.paragraphs[0]
            paragraph.text = lines[0] if lines else ''
        else:
            paragraph = cell.add_paragraph(lines[0] if lines else '')
        paragraph.alignment = alignment
        paragraph.paragraph_format.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(0)
        if bold:
            for run in paragraph.runs:
                run.bold = True
        for extra in lines[1:]:
            p = cell.add_paragraph()
            p.text = extra
            p.alignment = alignment
            p.paragraph_format.alignment = alignment
            p.paragraph_format.space_after = Pt(0)
            if bold:
                for run in p.runs:
                    run.bold = True
        try:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        except Exception:
            pass

    doc = Document()

    # increase bottom margin slightly so footer watermark stays visible below content
    try:
        section = doc.sections[0]
        # set margins to match PDF defaults (approx 15mm left/right, 30mm top, 22mm bottom)
        from docx.shared import Inches
        section.left_margin = Inches(15/25.4)
        section.right_margin = Inches(15/25.4)
        section.top_margin = Inches(30/25.4)
        # ensure header has some distance so inserted image is visible and doesn't overlap body
        try:
            section.header_distance = Inches(12/25.4)
        except Exception:
            pass
        section.bottom_margin = Inches(22/25.4)
    except Exception:
        section = None

    # use the original image at full opacity for DOCX background (not faded)
    # Prefer inserting directly from the image file path for reliability
    try:
        section = doc.sections[0]
        header = section.header
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        page_width = section.page_width - section.left_margin - section.right_margin
        wp = get_watermark_image_path()
        if wp is not None:
            try:
                run.add_picture(str(wp), width=page_width)
            except Exception:
                # fallback to stream-based insertion (keeps previous behavior)
                watermark_stream = create_watermark_image_bytes(opacity=1.0)
                if watermark_stream is not None:
                    try:
                        import tempfile, os
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                        tmp.write(watermark_stream.getvalue())
                        tmp.close()
                        run.add_picture(tmp.name, width=page_width)
                    finally:
                        try:
                            os.unlink(tmp.name)
                        except Exception:
                            pass
        else:
            watermark_stream = create_watermark_image_bytes(opacity=1.0)
            if watermark_stream is not None:
                try:
                    import tempfile, os
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    tmp.write(watermark_stream.getvalue())
                    tmp.close()
                    run.add_picture(tmp.name, width=page_width)
                finally:
                    try:
                        os.unlink(tmp.name)
                    except Exception:
                        pass
    except Exception:
        pass

    # add blank paragraphs before the title for spacing
    doc.add_paragraph('')
    doc.add_paragraph('')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run(model['title'])
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(8)
    # add one blank paragraph to move header down one line
    doc.add_paragraph('')

    # footer with updated date on every page
    try:
        if section is not None:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = model.get('updated_text', '')
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for run in footer_para.runs:
                run.font.size = Pt(9)
    except Exception:
        pass

    # DADOS GERAIS as a two-column table (labels left, values right)
    doc.add_heading('DADOS GERAIS', level=1)
    general_table = doc.add_table(rows=0, cols=2)
    general_table.style = 'Table Grid'
    general_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    general_table.autofit = False
    # compute column widths based on page usable width
    try:
        from docx.shared import Inches
        page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        general_table.columns[0].width = int(page_width * 0.62)
        general_table.columns[1].width = int(page_width * 0.38)
    except Exception:
        # fallback to approximate widths
        general_table.columns[0].width = Inches(4.5)
        general_table.columns[1].width = Inches(2.8)
    for label, value in model['general']:
        row_cells = general_table.add_row().cells
        # label as left cell (keep bold in visual but not a header)
        set_cell_text(row_cells[0], label)
        # make the label bold
        try:
            for run in row_cells[0].paragraphs[0].runs:
                run.bold = True
        except Exception:
            pass
        # special-case long 'Famílias' value to insert paragraph breaks
        if 'Família' in label or 'Famílias' in label:
            pieces = wrap_text_for_cell(value, max_chars=60)
            set_cell_text(row_cells[1], '\n'.join(pieces))
        else:
            set_cell_text(row_cells[1], value)

    doc.add_heading('DADOS DAS OBRAS', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    try:
        page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        table.columns[0].width = int(page_width * 0.62)
        table.columns[1].width = int(page_width * 0.19)
        table.columns[2].width = int(page_width * 0.19)
    except Exception:
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(1.4)
        table.columns[2].width = Inches(1.4)

    # add header row for DADOS DAS OBRAS
    header_cells = table.rows[0].cells
    set_cell_text(header_cells[0], 'Status', alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
    set_cell_text(header_cells[1], 'Quantidade', alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
    set_cell_text(header_cells[2], 'Valor (R$)', alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
    # header alignment: Status left, Quantidade and Valor centered
    try:
        for p in header_cells[0].paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for p in header_cells[1].paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for p in header_cells[2].paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        pass

    for row in model['summary']:
        row_cells = table.add_row().cells
        total_row = row['status'].strip().upper() == 'TOTAL'
        set_cell_text(row_cells[0], row['status'], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=total_row)
        set_cell_text(row_cells[1], str(row['quantidade']), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=total_row)
        set_cell_text(row_cells[2], clean_value(row['valor']), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=total_row)
        # align data: Status left, Quantidade center, Valor center
        try:
            for p in row_cells[0].paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for p in row_cells[1].paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for p in row_cells[2].paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            pass
        if total_row:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph('')
    doc.add_heading('DETALHAMENTO DAS OBRAS', level=1)
    for group in model['descricao']:
        for details in group['details']:
            para = doc.add_paragraph()
            try:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except Exception:
                pass
            status_run = para.add_run(f"{details['status']} - ")
            status_run.bold = True
            initiative_run = para.add_run(details['initiative'])
            initiative_run.bold = True
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), build_status_highlight_color(details['status']))
            status_run._r.get_or_add_rPr().append(shading)
            for label, value in details['rows']:
                row_para = doc.add_paragraph(style='List Bullet 2')
                try:
                    row_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                except Exception:
                    pass
                row_para.add_run(f"{label}: {value}")
            doc.add_paragraph('')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_odt_bytes(model):
    from odf.opendocument import OpenDocumentText
    from odf.style import Style, TextProperties, ParagraphProperties, TableColumnProperties, GraphicProperties
    from odf.draw import Frame, Image as DrawImage
    from odf.text import H, P, Span
    from odf.table import Table, TableRow, TableCell, TableColumn

    doc = OpenDocumentText()
    watermark_stream = create_watermark_image_bytes(opacity=0.15)
    if watermark_stream is not None:
        image_name = 'Watermark.png'
        doc.addPictureFromString(watermark_stream.getvalue(), 'image/png')
        watermark_style = Style(name='WatermarkGraphic', family='graphic')
        watermark_style.addElement(GraphicProperties(opacity='50%', wrap='none'))
        doc.styles.addElement(watermark_style)
        frame = Frame(stylename=watermark_style, anchortype='page', x='0cm', y='0cm', width='21cm', height='29.7cm', layer='background')
        draw_image = DrawImage(href='Pictures/' + image_name, type='simple', show='embed', actuate='onLoad')
        frame.addElement(draw_image)
        doc.text.addElement(frame)

    doc.text.addElement(P(text=''))
    doc.text.addElement(P(text=''))
    doc.text.addElement(H(outlinelevel='1', text=model['title']))
    doc.text.addElement(P(text=''))
    doc.text.addElement(P(text=''))
    doc.text.addElement(P(text=''))
    doc.text.addElement(H(outlinelevel='2', text='DADOS GERAIS'))

    bold_style = Style(name='BoldText', family='text')
    bold_style.addElement(TextProperties(fontweight='bold'))
    doc.styles.addElement(bold_style)

    # Create a two-column table for general data
    gen_col1 = Style(name='GeneralCol1', family='table-column')
    gen_col1.addElement(TableColumnProperties(columnwidth='9cm'))
    doc.styles.addElement(gen_col1)
    gen_col2 = Style(name='GeneralCol2', family='table-column')
    gen_col2.addElement(TableColumnProperties(columnwidth='8cm'))
    doc.styles.addElement(gen_col2)

    gen_table = Table()
    gen_table.addElement(TableColumn(stylename=gen_col1))
    gen_table.addElement(TableColumn(stylename=gen_col2))
    for label, value in model['general']:
        row = TableRow()
        cell = TableCell()
        cell.addElement(P(stylename=bold_style, text=f"{label}"))
        row.addElement(cell)
        cell = TableCell()
        cell.addElement(P(text=value))
        row.addElement(cell)
        gen_table.addElement(row)
    doc.text.addElement(gen_table)
    doc.text.addElement(P(text=''))
    doc.text.addElement(H(outlinelevel='2', text='DADOS DAS OBRAS'))

    col1_style = Style(name='ObraCol1', family='table-column')
    col1_style.addElement(TableColumnProperties(columnwidth='9cm'))
    doc.styles.addElement(col1_style)
    col2_style = Style(name='ObraCol2', family='table-column')
    col2_style.addElement(TableColumnProperties(columnwidth='4cm'))
    doc.styles.addElement(col2_style)
    col3_style = Style(name='ObraCol3', family='table-column')
    col3_style.addElement(TableColumnProperties(columnwidth='4cm'))
    doc.styles.addElement(col3_style)

    bold_style = Style(name='BoldText', family='text')
    bold_style.addElement(TextProperties(fontweight='bold'))
    doc.styles.addElement(bold_style)

    # paragraph style for centered text
    center_para = Style(name='CenterPara', family='paragraph')
    center_para.addElement(ParagraphProperties(textalign='center'))
    doc.styles.addElement(center_para)

    # paragraph style for centered bold text (used for headers)
    center_bold = Style(name='CenterBold', family='paragraph')
    center_bold.addElement(ParagraphProperties(textalign='center'))
    center_bold.addElement(TextProperties(fontweight='bold'))
    doc.styles.addElement(center_bold)

    initiative_style = Style(name='InitiativeBold', family='text')
    initiative_style.addElement(TextProperties(fontweight='bold'))
    doc.styles.addElement(initiative_style)

    table = Table()
    table.addElement(TableColumn(stylename=col1_style))
    table.addElement(TableColumn(stylename=col2_style))
    table.addElement(TableColumn(stylename=col3_style))
    # add header row for ODT DADOS DAS OBRAS
    header_row = TableRow()
    hcell = TableCell()
    hcell.addElement(P(stylename=bold_style, text='Status'))
    header_row.addElement(hcell)
    hcell = TableCell()
    hcell.addElement(P(stylename=center_bold, text='Quantidade'))
    header_row.addElement(hcell)
    hcell = TableCell()
    hcell.addElement(P(stylename=center_bold, text='Valor (R$)'))
    header_row.addElement(hcell)
    table.addElement(header_row)

    for row in model['summary']:
        row_cells = TableRow()
        total_row = row['status'].strip().upper() == 'TOTAL'

        cell = TableCell()
        if total_row:
            cell.addElement(P(stylename=bold_style, text=row['status']))
        else:
            cell.addElement(P(text=row['status']))
        row_cells.addElement(cell)

        cell = TableCell()
        if total_row:
            cell.addElement(P(stylename=center_bold, text=str(row['quantidade'])))
        else:
            cell.addElement(P(stylename=center_para, text=str(row['quantidade'])))
        row_cells.addElement(cell)

        cell = TableCell()
        if total_row:
            cell.addElement(P(stylename=center_bold, text=row['valor'].replace('R$', '').strip()))
        else:
            cell.addElement(P(stylename=center_para, text=row['valor'].replace('R$', '').strip()))
        row_cells.addElement(cell)

        table.addElement(row_cells)
    doc.text.addElement(table)

    doc.text.addElement(P(text=''))
    doc.text.addElement(H(outlinelevel='2', text='DETALHAMENTO DAS OBRAS'))

    style_cache = {}
    def get_status_style_name(status):
        key = status.lower().replace(' ', '_')
        if key in style_cache:
            return style_cache[key]
        style_name = f"StatusHighlight_{key}"
        style = Style(name=style_name, family="text")
        style.addElement(TextProperties(color="#000000", backgroundcolor=f"#{build_status_highlight_color(status)}", fontweight="bold"))
        doc.styles.addElement(style)
        style_cache[key] = style_name
        return style_name

    for group in model['descricao']:
        for details in group['details']:
            para = P()
            status_style_name = get_status_style_name(details['status'])
            status_span = Span(text=f"{details['status']} - ", stylename=status_style_name)
            initiative_span = Span(text=details['initiative'], stylename='InitiativeBold')
            para.addElement(status_span)
            para.addElement(initiative_span)
            doc.text.addElement(para)
            for label, value in details['rows']:
                doc.text.addElement(P(text=f'    • {label}: {value}'))
            doc.text.addElement(P(text=''))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_document_model_from_text(document_text):
    lines = document_text.splitlines()
    model = {'title': '', 'general': [], 'summary': [], 'descricao': []}
    section = None
    current_group = None
    current_details = []

    for line in lines:
        stripped = line.strip()
        if stripped == '':
            continue
        if stripped == 'DADOS GERAIS':
            section = 'general'
            continue
        if stripped == 'DADOS DAS OBRAS':
            section = 'summary'
            continue
        if stripped == 'DETALHAMENTO DAS OBRAS':
            section = 'descricao'
            current_detail = None
            if not model['descricao']:
                model['descricao'].append({'status': 'DETALHAMENTO DAS OBRAS', 'details': []})
            continue
        if section is None:
            model['title'] = stripped
            continue
        if section == 'general':
            line_value = stripped[2:].strip() if stripped.startswith('• ') else stripped
            if ':' in line_value:
                label, value = line_value.split(':', 1)
                model['general'].append((label.strip(), value.strip()))
        elif section == 'summary':
            line_value = stripped[2:].strip() if stripped.startswith('• ') else stripped
            if ':' in line_value and '— Valor (R$):' in line_value:
                status_part, rest = line_value.split(':', 1)
                qtd_part, valor_part = rest.split('— Valor (R$):')
                quantidade = qtd_part.replace('Quantidade', '').strip()
                model['summary'].append({'status': status_part.strip(), 'quantidade': int(quantidade), 'valor': valor_part.strip()})
        elif section == 'descricao':
            if ' - ' in stripped and not stripped.startswith('• '):
                status_part, initiative_part = stripped.split(' - ', 1)
                current_detail = {
                    'status': status_part.strip(),
                    'initiative': initiative_part.strip(),
                    'status_color': build_status_highlight_color(status_part.strip()),
                    'rows': []
                }
                model['descricao'][0]['details'].append(current_detail)
            elif stripped.startswith('• '):
                if current_detail is None:
                    continue
                if ':' in stripped[2:]:
                    label, value = stripped[2:].split(':', 1)
                    current_detail['rows'].append((label.strip(), value.strip()))
                else:
                    current_detail['rows'].append(('Descrição', stripped[2:].strip()))
            else:
                if current_detail is None:
                    continue
                current_detail['rows'].append(('Descrição', stripped))
    return model


def is_reportlab_available():
    try:
        import reportlab  # noqa: F401
        return True
    except ImportError:
        return False


def remove_trailing_blank_pages(pdf_bytes):
    if not pdf_bytes:
        return pdf_bytes

    try:
        from PyPDF2 import PdfReader, PdfWriter
        from io import BytesIO as _BI
    except ImportError:
        try:
            from pypdf import PdfReader, PdfWriter
            from io import BytesIO as _BI
        except ImportError:
            return pdf_bytes

    try:
        reader = PdfReader(_BI(pdf_bytes))
        if len(reader.pages) <= 1:
            return pdf_bytes

        pages_to_keep = list(reader.pages)
        while len(pages_to_keep) > 1:
            last_page = pages_to_keep[-1]
            try:
                last_text = last_page.extract_text() or ''
            except Exception:
                last_text = ''

            normalized_text = re.sub(r'\s+', ' ', (last_text or '')).strip()
            text_without_punctuation = re.sub(r'[^A-Za-zÀ-ÿ0-9]', '', normalized_text)
            page_has_meaningful_text = bool(text_without_punctuation) and len(text_without_punctuation) > 8

            if page_has_meaningful_text:
                break

            if normalized_text and not re.fullmatch(r'\d+', normalized_text):
                if len(normalized_text.split()) <= 2 and len(normalized_text) <= 24:
                    pages_to_keep = pages_to_keep[:-1]
                    continue

            # If the page has no meaningful extracted text, it is treated as blank.
            # This avoids keeping a final empty page that only carries the footer or page number.
            pages_to_keep = pages_to_keep[:-1]

        if len(pages_to_keep) == len(reader.pages):
            return pdf_bytes

        writer = PdfWriter()
        for page in pages_to_keep:
            writer.add_page(page)

        out_b = _BI()
        writer.write(out_b)
        return out_b.getvalue()
    except Exception:
        return pdf_bytes


def create_pdf_bytes(model):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_RIGHT, TA_LEFT
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, NextPageTemplate, Paragraph, Spacer, Table, TableStyle

    buffer = BytesIO()
    # increase bottom margin so footer watermark remains visible
    # use two page templates: one for the first page (smaller top margin)
    # and one for later pages (larger top margin) so that the first-page
    # title stays higher while subsequent pages start lower and avoid header overlap.
    left_margin = 15 * mm
    right_margin = 15 * mm
    bottom_margin = 22 * mm
    top_margin_first = 30 * mm
    top_margin_later = 40 * mm

    page_width, page_height = A4
    frame_width = page_width - left_margin - right_margin
    frame_height_first = page_height - top_margin_first - bottom_margin
    frame_height_later = page_height - top_margin_later - bottom_margin

    doc = BaseDocTemplate(buffer, pagesize=A4,
                          leftMargin=left_margin, rightMargin=right_margin,
                          topMargin=top_margin_first, bottomMargin=bottom_margin)

    first_frame = Frame(left_margin, bottom_margin, frame_width, frame_height_first, id='first')
    later_frame = Frame(left_margin, bottom_margin, frame_width, frame_height_later, id='later')

    first_template = PageTemplate(id='First', frames=[first_frame], onPage=None)
    later_template = PageTemplate(id='Later', frames=[later_frame], onPage=None)
    doc.addPageTemplates([first_template, later_template])
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleCenter', parent=styles['Title'], alignment=1, spaceBefore=12, spaceAfter=8))
    styles.add(ParagraphStyle(name='Section', parent=styles['Heading2'], spaceAfter=6, spaceBefore=12, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name='SectionCompact', parent=styles['Heading2'], spaceAfter=4, spaceBefore=4, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name='NormalIndent', parent=styles['BodyText'], leftIndent=5*mm, spaceAfter=2, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name='GeneralLabel', parent=styles['BodyText'], leftIndent=0, spaceAfter=2, alignment=TA_JUSTIFY, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='CenterLabel', parent=styles['BodyText'], alignment=1, fontName='Helvetica-Bold', spaceAfter=2))
    styles.add(ParagraphStyle(name='CenterData', parent=styles['BodyText'], alignment=1, leftIndent=0, rightIndent=0, spaceAfter=2))
    styles.add(ParagraphStyle(name='LeftData', parent=styles['BodyText'], alignment=TA_LEFT, leftIndent=0, rightIndent=0, spaceAfter=2))
    styles.add(ParagraphStyle(name='CenterDataBold', parent=styles['BodyText'], alignment=1, leftIndent=0, rightIndent=0, spaceAfter=2, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='LeftDataBold', parent=styles['BodyText'], alignment=TA_LEFT, leftIndent=0, rightIndent=0, spaceAfter=2, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='BulletIndent', parent=styles['BodyText'], leftIndent=10*mm, bulletIndent=5*mm, spaceAfter=2, alignment=TA_JUSTIFY))

    # use original image at full opacity for the PDF so it's sharp (not faded)
    watermark_stream = create_watermark_image_bytes(opacity=1.0)
    watermark_image = ImageReader(watermark_stream) if watermark_stream is not None else None

    story = []
    # keep first-page spacing (title appears higher) but tightened to fit more content
    story.append(Spacer(1, 4*mm))
    story.append(Paragraph(model['title'], styles['TitleCenter']))
    # reduce space between title and DADOS GERAIS so more content fits on final page
    story.append(Spacer(1, 4*mm))
    # extra small spacer to move header down one line for watermark visibility
    story.append(Spacer(1, 2*mm))
    # ensure subsequent pages use the 'Later' template (consumed on page 1)
    story.append(NextPageTemplate('Later'))

    story.append(Paragraph('DADOS GERAIS', styles['Section']))
    general_data = []
    for label, value in model['general']:
        # For long 'Famílias' field, insert manual breaks for readability
        if 'Família' in label or 'Famílias' in label:
            pieces = []
            words = str(value).split()
            line = ''
            for w in words:
                if len(line) + 1 + len(w) <= 60:
                    line = (line + ' ' + w).strip()
                else:
                    pieces.append(line)
                    line = w
            if line:
                pieces.append(line)
            value_text = '<br/>'.join(pieces)
        else:
            value_text = value
        general_data.append([
            Paragraph(label, styles['GeneralLabel']),
            Paragraph(value_text, styles['NormalIndent'])
        ])
    table = Table(general_data, colWidths=[90*mm, 85*mm])
    table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('WORDWRAP', (0, 0), (-1, -1), 'CJK'),
    ]))
    story.append(table)
    story.append(Spacer(1, 6*mm))

    story.append(Paragraph('DADOS DAS OBRAS', styles['SectionCompact']))
    summary_data_wrapped = [[
        Paragraph('Status', styles['GeneralLabel']),
        Paragraph('Quantidade', styles['CenterLabel']),
        Paragraph('Valor (R$)', styles['CenterLabel'])
    ]]
    for row in model['summary']:
        total_row = str(row['status']).strip().upper() == 'TOTAL'
        summary_data_wrapped.append([
            Paragraph(row['status'], styles['LeftDataBold'] if total_row else styles['LeftData']),
            Paragraph(str(row['quantidade']), styles['CenterDataBold'] if total_row else styles['CenterData']),
            Paragraph(row['valor'].replace('R$', '').strip(), styles['CenterDataBold'] if total_row else styles['CenterData'])
        ])
    table = Table(summary_data_wrapped, colWidths=[95*mm, 40*mm, 40*mm])
    last_row = len(summary_data_wrapped) - 1
    table_style = [
        ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (2, -1), 'CENTER'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]
    if last_row >= 1:
        # last_row index corresponds to last data row; make last row bold
        table_style.append(('FONTNAME', (0, last_row), (-1, last_row), 'Helvetica-Bold'))
    table.setStyle(TableStyle(table_style))
    story.append(table)
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph('DETALHAMENTO DAS OBRAS', styles['SectionCompact']))
    status_style_registry = {}
    detail_items = []
    for group in model['descricao']:
        detail_items.extend(group['details'])

    for index, details in enumerate(detail_items):
        status_color = colors.HexColor(f"#{build_status_highlight_color(details['status'])}")
        key = details['status'].lower().replace(' ', '_')
        style_name = f'StatusDetail_{key}'
        if style_name not in status_style_registry:
            styles.add(ParagraphStyle(
                name=style_name,
                parent=styles['Heading3'],
                textColor=colors.black,
                spaceAfter=2,
                leftIndent=0,
                alignment=TA_JUSTIFY
            ))
            status_style_registry[style_name] = True
        status_text = f"<font backColor='{status_color}' color='black'><b>{details['status']}</b></font> - <b>{details['initiative']}</b>"
        story.append(Paragraph(status_text, styles[style_name]))
        for label, value in details['rows']:
            story.append(Paragraph(f'• {label}: {value}', styles['NormalIndent']))
        if index < len(detail_items) - 1:
            story.append(Spacer(1, 1*mm))

    def footer(canvas, doc):
        canvas.saveState()
        page_num = canvas.getPageNumber()
        canvas.setFont('Helvetica', 9)
        updated_text = model.get('updated_text', '')
        if updated_text:
            canvas.drawRightString(doc.pagesize[0] - 15*mm, 16*mm, updated_text)
        canvas.drawString(15*mm, 10*mm, str(page_num))
        canvas.restoreState()

    def draw_page(canvas, doc):
        canvas.saveState()
        if watermark_image is not None:
            # draw full-page image preserving aspect ratio to keep it sharp
            try:
                canvas.drawImage(watermark_image, 0, 0, width=doc.pagesize[0], height=doc.pagesize[1], mask='auto', preserveAspectRatio=True)
            except TypeError:
                # older reportlab versions may not support preserveAspectRatio
                canvas.drawImage(watermark_image, 0, 0, width=doc.pagesize[0], height=doc.pagesize[1], mask='auto')
        footer(canvas, doc)
        canvas.restoreState()

    # attach header/footer drawing to both page templates
    # set the onPage function for the templates so header/footer are drawn
    for tpl in doc.pageTemplates:
        tpl.onPage = draw_page

    doc.build(story)

    # post-process: remove trailing pages that are effectively blank
    # to avoid an extra page at the end of the generated PDF.
    return remove_trailing_blank_pages(buffer.getvalue())


def generate_report(dataframes, informar_option='Todos'):
    """Generate the complete report data"""
    df_segreg = dataframes.get('Segreg por município', dataframes.get('Investimentos', pd.DataFrame()))
    df_segreg = filter_by_informar_no_relatorio(df_segreg, informar_option)

    df_volumes = dataframes.get('Volumes por município', dataframes.get('Volumes por municipio', pd.DataFrame()))
    df_abastecimento = dataframes.get('Abastecimento', pd.DataFrame())
    df_populacao = dataframes.get('População por município', dataframes.get('Populacao por município', pd.DataFrame()))
    df_tv_tspe = dataframes.get('TV+TSPE', dataframes.get('TV + TSPE', pd.DataFrame()))
    
    dados_gerais = extract_dados_gerais(df_segreg, df_volumes, df_abastecimento, df_populacao, df_tv_tspe)
    dados_obras = extract_dados_obras(df_segreg)
    descricao_obras = extract_descricao_obras(df_segreg)
    
    return dados_gerais, dados_obras, descricao_obras

def display_dados_gerais(dados_gerais):
    """Display general data in the Streamlit interface"""
    st.header("📋 DADOS GERAIS")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Sistemas de Abastecimento:**")
        st.write(normalize_missing_value(dados_gerais.get('sistema_abastecimento', '[Informação não disponível]')))
        
        st.markdown("**Produção total de água:**")
        st.write(normalize_missing_value(dados_gerais.get('producao_total_agua', '[Informação não disponível]')))
        
        st.markdown("**Calendário médio de abastecimento:**")
        st.write(normalize_missing_value(dados_gerais.get('calendario_medio', '[Informação não disponível]')))
    
    with col2:
        calendario_pos_obras = dados_gerais.get('calendario_pos_obras', None)
        if has_effective_value(calendario_pos_obras):
            st.markdown("**Calendário médio pós obras:**")
            st.write(normalize_missing_value(calendario_pos_obras, '[Informação não disponível]'))
        
        st.markdown("**Famílias beneficiadas com a Tarifa Social Pernambucana e a Tarifa de Vulneráveis:**")
        st.write(format_integer_pt_br(dados_gerais.get('familias_beneficiadas', '[Informação não disponível]')))
        
        st.markdown("**População:**")
        st.write(normalize_missing_value(dados_gerais.get('populacao', '[Informação não disponível]')))

def summarize_obras_for_report(selected_obras_by_status=None, dados_obras=None):
    """Summarize obras using the explicitly selected items when available, otherwise fall back to the full dataset."""
    summary_data = []
    total_quantidade = 0
    total_valor = 0

    if selected_obras_by_status is not None:
        for status in STATUS_ORDER:
            obras = selected_obras_by_status.get(status, []) or []
            if not obras:
                continue

            qtd = len(obras)
            valor = sum(
                get_investimento_municipio(obra, 0)
                for obra in obras
            )
            summary_data.append({
                'Status': status,
                'Quantidade': qtd,
                'Valor (R$)': valor
            })
            total_quantidade += qtd
            total_valor += valor
    elif dados_obras:
        for status in STATUS_ORDER:
            if status in dados_obras:
                qtd = dados_obras[status]['quantidade']
                if qtd <= 0:
                    continue
                valor = dados_obras[status]['valor']
                summary_data.append({
                    'Status': status,
                    'Quantidade': qtd,
                    'Valor (R$)': valor
                })
                total_quantidade += qtd
                total_valor += valor

    summary_data.append({
        'Status': 'TOTAL',
        'Quantidade': total_quantidade,
        'Valor (R$)': total_valor
    })

    return summary_data


def display_dados_obras(dados_obras, selected_obras_by_status=None):
    """Display obra data summary"""
    st.header("📊 DADOS DAS OBRAS")

    summary_data = summarize_obras_for_report(
        selected_obras_by_status=selected_obras_by_status,
        dados_obras=dados_obras,
    )
    
    df_summary = pd.DataFrame(summary_data)
    
    # Format summary values in pt_BR with currency rounding
    df_summary['Valor (R$)'] = df_summary['Valor (R$)'].apply(format_currency_rounded)
    
    def highlight_total(row):
        if str(row['Status']).strip().upper() == 'TOTAL':
            return ['font-weight: bold'] * len(row)
        return [''] * len(row)

    styled_summary = (
        df_summary.style
        .set_properties(subset=['Status'], **{'text-align': 'left'})
        .apply(highlight_total, axis=1)
    )

    st.dataframe(styled_summary, use_container_width=True, hide_index=True)

def escape_markdown_dollars(text):
    return text.replace('$', '\\$') if isinstance(text, str) else text


def display_descricao_obras(descricao_obras, df_segreg):
    """Display detailed description of obras"""
    st.header("📝 DESCRIÇÃO DAS OBRAS")
    
    # Get status mapping for display text
    status_display = {
        'CONCLUÍDO': 'CONCLUÍDA',
        'EM ANDAMENTO': 'EM ANDAMENTO',
        'A INICIAR': 'A INICIAR',
        'EM LICITAÇÃO': 'EM LICITAÇÃO',
        'A LICITAR': 'A LICITAR',
        'A FAZER': 'A FAZER'
    }
    
    selected_obras_by_status = {}

    for status in STATUS_ORDER:
        if status in descricao_obras and descricao_obras[status]:
            st.subheader(f"**{status_display.get(status, status)}**")
            
            obras = descricao_obras[status]
            selected_list = []
            for idx, obra in enumerate(obras):
                nome_iniciativa = normalize_missing_value(get_first_key(obra, ['Nome da Iniciativa', 'Nome da iniciativa'], '[Informação não disponível]'))
                pop_beneficiada = format_population_value(get_first_key(obra, ['População Beneficiada', 'Populacao Beneficiada'], '[Informação não disponível]'))
                municipios = normalize_missing_value(get_first_key(obra, ['Todos os municípios beneficiados', 'Todos os municipios beneficiados'], '[Informação não disponível]'))
                
                investimento = get_investimento_municipio(obra, 0)
                investimento_total = get_numeric_value(
                    obra,
                    ['Total do Investimento da iniciativa R$', 'Total do Investimento da iniciativa', 'Total do Investimento R$', 'Total do Investimento R$ por município'],
                    investimento,
                )
                fonte = normalize_missing_value(get_first_key(obra, ['Fonte de Recurso', 'Fonte de recurso'], 'Governo de Pernambuco e Compesa'))
                prazo = get_first_key(obra, ['Prazo de Conclusão', 'Prazo de Conclusao'], '')
                data_inicio = get_first_key(obra, ['Data Início', 'Data Inicio'], '')
                proxima_etapa = get_first_key(obra, ['Próxima Etapa', 'Proxima Etapa'], '')
                data_proxima_etapa = get_first_key(obra, ['Data Próxima Etapa', 'Data Proxima Etapa'], '')

                # Status-specific text
                if status == 'CONCLUÍDO':
                    status_text = f"**CONCLUÍDA** - {nome_iniciativa}. {pop_beneficiada} pessoas beneficiadas dos municípios de {municipios}."
                    prazo_text = f"\n\nConclusão: {format_month_year_pt_br(prazo)}" if prazo else ""
                    proxima_etapa_text = ""
                elif status == 'EM ANDAMENTO':
                    status_text = f"**EM ANDAMENTO** - {nome_iniciativa}. {pop_beneficiada} pessoas beneficiadas dos municípios de {municipios}."
                    prazo_text = f"\n\nPrevisão de término: {format_month_year_pt_br(prazo)}" if prazo else ""
                    proxima_etapa_text = ""
                elif status == 'EM LICITAÇÃO':
                    status_text = f"**EM LICITAÇÃO** - {nome_iniciativa}. {pop_beneficiada} pessoas beneficiadas dos municípios de {municipios}."
                    prazo_text = ""
                    if data_inicio:
                        prazo_text += f"\n\nPrevisão de início: {format_month_year_pt_br(data_inicio)}"
                    if prazo:
                        prazo_text += f"\nPrevisão de término: {format_month_year_pt_br(prazo)}"
                    proxima_etapa_text = ""
                elif status in ('A INICIAR', 'A LICITAR', 'A FAZER'):
                    status_text = f"**{status}** - {nome_iniciativa}. {pop_beneficiada} pessoas beneficiadas dos municípios de {municipios}."
                    prazo_text = ""
                    formatted_proxima = format_date_pt_br(proxima_etapa)
                    if status == 'A LICITAR' or status == 'A INICIAR':
                        formatted_data = format_month_year_pt_br(data_proxima_etapa)
                    else:
                        formatted_data = format_date_pt_br(data_proxima_etapa)
                    data_label = 'Previsão de emissão de ordem de serviço:' if status in ('A INICIAR', 'A FAZER') else 'Previsão de publicação do edital de licitação:'
                    if formatted_proxima and formatted_data:
                        proxima_etapa_text = f"\n\nPróxima etapa: {formatted_proxima}. {data_label} {formatted_data}"
                    elif formatted_proxima:
                        proxima_etapa_text = f"\n\nPróxima etapa: {formatted_proxima}."
                    elif formatted_data:
                        proxima_etapa_text = f"\n\n{data_label} {formatted_data}"
                    else:
                        proxima_etapa_text = ""
                else:
                    status_text = f"**{status}** - {nome_iniciativa}. {pop_beneficiada} pessoas beneficiadas dos municípios de {municipios}."
                    prazo_text = ""
                    proxima_etapa_text = ""
                
                invest_text = f"Investimento: {format_currency_rounded(investimento)} ({fonte})"
                invest_text += f"\nTotal do Investimento: {format_currency_rounded(investimento_total)}"
                
                full_text = status_text + f"\n\n{invest_text}"
                full_text += prazo_text
                full_text += proxima_etapa_text

                cols = st.columns([0.94, 0.06])
                with cols[0]:
                    st.markdown(escape_markdown_dollars(full_text))
                with cols[1]:
                    include = st.checkbox("Incluir", value=True, key=f"include_{status}_{idx}")

                if include:
                    selected_list.append(obra)

                st.divider()

            if selected_list:
                selected_obras_by_status[status] = selected_list

    return selected_obras_by_status

def main():
    st.title("🏗️ Gerador de Nota Técnica")
    st.markdown("---")
    
    dataframes = None
    supabase_configured = False

    if db_supabase:
        creds = db_supabase.get_supabase_credentials()
        if creds.get("db_url"):
            supabase_configured = True

    # PAINEL LATERAL (ADMINISTRADOR COM SENHA)
    if db_supabase:
        with st.sidebar.expander("⚙️ Painel de Administração", expanded=False):
            if supabase_configured:
                st.success("🟢 Supabase Conectado!")
                if st.button("🔌 Testar Conexão"):
                    ok, msg = db_supabase.test_supabase_connection()
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)

                st.markdown("---")
                st.subheader("🔑 Acesso do Administrador")
                
                # Obtém a senha do st.secrets ou usa a padrão 'compesa2026'
                admin_pass_secret = None
                if hasattr(st, "secrets"):
                    admin_pass_secret = st.secrets.get("ADMIN_PASSWORD", "compesa2026")
                if not admin_pass_secret:
                    admin_pass_secret = "compesa2026"

                entered_pass = st.text_input("Digite a Senha de Administrador:", type="password", key="admin_password_input")
                
                if entered_pass == admin_pass_secret:
                    st.success("🔓 Modo Administrador Ativado")
                    st.subheader("📤 Atualizar Planilha no Banco")
                    admin_file = st.file_uploader(
                        "Envie a nova versão do Excel",
                        type=['xlsx', 'xls'],
                        key="admin_excel_upload"
                    )
                    if admin_file is not None:
                        admin_dfs = load_excel_file(admin_file)
                        if st.button("☁️ Substituir Base no Supabase"):
                            with st.spinner("Atualizando Supabase..."):
                                if db_supabase.upload_excel_dict_to_supabase(admin_dfs):
                                    st.success("✅ Nova planilha gravada com sucesso no Supabase!")
                                    st.cache_data.clear()
                                    st.rerun()
                elif entered_pass:
                    st.error("❌ Senha incorreta.")
                else:
                    st.info("🔒 Digite a senha para desbloquear a área de atualização de planilhas.")
            else:
                st.info("Supabase pendente. Configure em `.streamlit/secrets.toml` ou `.env`.")


    # FLUXO PRINCIPAL DA TELA (USUÁRIO FINAL)
    if supabase_configured:
        col_title, col_reload = st.columns([0.75, 0.25])
        with col_reload:
            if st.button("🔄 Recarregar Base"):
                st.cache_data.clear()
                st.rerun()

        with st.spinner("⚡ Carregando base de dados..."):
            dataframes = db_supabase.load_all_sheets_from_supabase()

        if not dataframes:
            st.warning("⚠️ Nenhuma base encontrada no Supabase. Abra o menu '⚙️ Administração' na barra lateral para enviar o primeiro arquivo Excel.")
    else:
        # Fallback caso o Supabase não esteja configurado
        uploaded_file = st.file_uploader(
            "📂 Carregue o arquivo Excel localmente",
            type=['xlsx', 'xls'],
            help="O arquivo deve conter as abas: 'Segreg por município', 'Volumes por município', 'Abastecimento', 'População por município'"
        )

        raw_excel_file = None
        upload_complete = False

        if 'downloaded_excel_bytes' in st.session_state and st.session_state['downloaded_excel_bytes']:
            raw_excel_file = st.session_state['downloaded_excel_bytes']
        elif uploaded_file is not None:
            raw_excel_file = uploaded_file
            upload_complete = True

        if uploaded_file is not None:
            current_name = getattr(uploaded_file, 'name', None) or None
            if st.session_state.get('last_uploaded_name') != current_name:
                st.session_state['last_uploaded_name'] = current_name
                st.session_state['municipios_loaded'] = False
                st.session_state['show_upload_message'] = True

        upload_msg_ph = st.empty()
        if upload_complete and not st.session_state.get('municipios_loaded', False) and st.session_state.get('show_upload_message', True):
            upload_msg_ph.warning("Upload completo! ✔ Aguarde até que as informações dos municípios sejam carregadas.")
        else:
            upload_msg_ph.empty()

        if raw_excel_file is not None:
            dataframes = load_excel_file(raw_excel_file)

    if dataframes:
        try:


            
            segreg_sheet = find_sheet_name(dataframes, ['Segreg por município', 'Segreg por municipio', 'Investimentos', 'Investimento'], required_columns=['Status', 'Status', 'Total do Investimento R$ por município', 'Total do Investimento'])
            volumes_sheet = find_sheet_name(dataframes, ['Volumes por município', 'Volumes por municipio', 'Volumes'], required_columns=['Vazão [L/s]', 'Vazao [L/s]', 'Vazão', 'Vazao'])
            abastecimento_sheet = find_sheet_name(dataframes, ['Abastecimento'], required_columns=['Sistema de Abastecimento', 'Calendário Médio Atual', 'Calendário Médio Pós Obras', 'População Total Residente SINISA 2024'])
            populacao_sheet = find_sheet_name(dataframes, ['População por município', 'Populacao por município', 'População por municipio', 'Populacao por municipio'], required_columns=['TV + TSPE (Econ)', 'População Beneficiada', 'Populacao Beneficiada'])
            tv_tspe_sheet = find_sheet_name(dataframes, ['TV+TSPE', 'TV + TSPE', 'TV+TSPE (Econ)', 'TV + TSPE (Econ)'], required_columns=['TV + TSPE (Econ)'])
            investimentos_sheet = find_sheet_name(dataframes, ['Investimentos', 'Investimento'], required_columns=['Município Principal', 'Municipio Principal'])

            resolved_dataframes = {
                'Segreg por município': dataframes.get(segreg_sheet, pd.DataFrame()) if segreg_sheet else pd.DataFrame(),
                'Volumes por município': dataframes.get(volumes_sheet, pd.DataFrame()) if volumes_sheet else pd.DataFrame(),
                'Abastecimento': dataframes.get(abastecimento_sheet, pd.DataFrame()) if abastecimento_sheet else pd.DataFrame(),
                'População por município': dataframes.get(populacao_sheet, pd.DataFrame()) if populacao_sheet else pd.DataFrame(),
                'TV+TSPE': dataframes.get(tv_tspe_sheet, pd.DataFrame()) if tv_tspe_sheet else pd.DataFrame(),
                'Investimentos': dataframes.get(investimentos_sheet, pd.DataFrame()) if investimentos_sheet else pd.DataFrame()
            }

            municipios = collect_municipios(resolved_dataframes)
            # mark as loaded so the upload confirmation message is hidden
            st.session_state['municipios_loaded'] = True
            # hide the explicit upload message now that municipio options will be shown
            st.session_state['show_upload_message'] = False
            try:
                upload_msg_ph.empty()
            except Exception:
                pass
            if not municipios:
                st.error("⚠️ Não foi possível identificar municípios válidos nas abas do arquivo. Verifique as colunas de município.")
                st.stop()

            municipio_options = ["Selecione um município"] + municipios
            selected_municipio = st.selectbox("📍 Escolha o município", municipio_options)
            if selected_municipio == "Selecione um município":
                st.warning("Selecione um município para gerar a Nota Técnica.")
                st.stop()

            # Filtro Informar no Relatório
            informar_option = st.selectbox(
                "Informar no Relatório",
                options=["Todos", "Sim", "Não"],
                index=0,
                key="informar_no_relatorio_filter"
            )

            filtered_dataframes = {
                'Segreg por município': filter_by_municipio(resolved_dataframes['Segreg por município'], selected_municipio, 'Segreg por município'),
                'Volumes por município': filter_by_municipio(resolved_dataframes['Volumes por município'], selected_municipio, 'Volumes por município'),
                'Abastecimento': filter_by_municipio(resolved_dataframes['Abastecimento'], selected_municipio, 'Abastecimento'),
                'População por município': filter_by_municipio(resolved_dataframes['População por município'], selected_municipio, 'População por município'),
                'TV+TSPE': filter_by_municipio(resolved_dataframes['TV+TSPE'], selected_municipio, 'TV+TSPE'),
                'Investimentos': filter_by_municipio(resolved_dataframes['Investimentos'], selected_municipio, 'Investimentos')
            }

            dados_gerais, dados_obras, descricao_obras = generate_report(filtered_dataframes, informar_option=informar_option)

            total_investment = sum(
                float(v.get('valor', 0) or 0)
                for v in dados_obras.values()
            )

            st.header(f"📝 Município de {selected_municipio}")
            st.markdown("---")

            df_segreg = filtered_dataframes.get('Segreg por município', pd.DataFrame())
            df_segreg = filter_by_informar_no_relatorio(df_segreg, informar_option)

            # Display sections
            display_dados_gerais(dados_gerais)
            
            st.markdown("---")

            st.header("📊 DADOS DAS OBRAS")
            if total_investment <= 0:
                st.markdown(f"""
                <div style="display:flex; justify-content:center; align-items:center; min-height:100px; margin-bottom:20px;">
                    <div style="border:1px solid rgba(0,0,0,0.15); border-radius:20px; padding:20px 30px; background-color:#fff9c4; text-align:center; max-width:980px;">
                        <span style="font-size:24px; font-weight:700; color:#111; line-height:1.2;">O Município {selected_municipio} não possui investimentos</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                # Criar dicionário vazio para não quebrar o modelo
                selected_descricao_obras = {}
            else:
                selected_descricao_obras = display_descricao_obras(descricao_obras, df_segreg)
                display_dados_obras(dados_obras, selected_descricao_obras)

            st.markdown("---")
            st.header("📥 Exportar documento")
            model = build_document_model(selected_descricao_obras, dados_gerais, dados_obras, selected_municipio)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            try:
                import reportlab
                pdf_bytes = create_pdf_bytes(model)
                st.download_button(
                    label="Baixar .pdf",
                    data=pdf_bytes,
                    file_name=f"{selected_municipio}_{timestamp}.pdf",
                    mime="application/pdf"
                )
            except ImportError:
                st.error("Instale 'reportlab' para exportar em .pdf.")
            except Exception as e:
                st.error(f"Erro ao gerar PDF: {e}")

            # try:
            #     odt_bytes = create_odt_bytes(model)
            #     st.download_button(
            #         label="Baixar .odt",
            #         data=odt_bytes,
            #         file_name=f"nota_tecnica_{selected_municipio}_{timestamp}.odt",
            #         mime="application/vnd.oasis.opendocument.text"
            #     )
            # except ImportError:
            #     st.warning("Instale 'odfpy' para exportar em .odt.")
        except Exception as e:
            st.error(f"❌ Erro ao processar o arquivo: {str(e)}")
            st.exception(e)
    else:
        st.info("👆 Carregue um arquivo Excel para começar a gerar o relatório.")
        
        # Show expected format
        # with st.expander("ℹ️ Formato esperado do arquivo"):
        #     st.markdown("""
        #     O arquivo Excel deve conter as sseguintes abas:
            
        #     **1. Segreg por município**
        #     - Colunas esperadas: 'Status', 'População Beneficiada', 'Todos os municípios beneficiados', 
        #       'Total do Investimento R$ por município', 'Total do Investimento R$', 'Fonte de Recurso', 
        #       'Prazo de Conclusão', 'Data Início'
            
        #     **2. Volumes por município**
        #     - Colunas esperadas: 'Vazão [L/s]'
            
        #     **3. Abastecimento**
        #     - Colunas esperadas: 'Calendário Médio Atual', 'Calendário Médio Pós Obras'
            
        #     **4. População por município**
        #     - Colunas esperadas: 'População', 'TV', 'TSPE (Econ)'
        #     """)

if __name__ == "__main__":
    main()